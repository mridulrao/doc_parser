# app.py
import io
import hashlib
import re
import tempfile
from pathlib import Path

import streamlit as st

from agent import (
    clean_document_content,
    annotated_document_to_html,
    run_form_chat_turn,
    annotated_to_plain,
)
from docs_to_pdf import convert_docx_to_pdf
from pypdf import PdfReader
from docx import Document

import asyncio
import threading
import time


def arun(coro):
    """Helper to run async coroutines from Streamlit (sync) code."""
    return asyncio.run(coro)

_INVALID_XML_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]")

def _sanitize_for_docx(text: str) -> str:
    if not isinstance(text, str):
        text = str(text)
    return _INVALID_XML_RE.sub("", text)

st.set_page_config(page_title="Doc Assist", page_icon="📄", layout="wide")
st.title("Doc Assist")

uploaded = st.file_uploader(
    "Choose a file",
    type=["pdf", "docx"],
    accept_multiple_files=False,
    help="Drag & drop or browse to upload a single document.",
)

if uploaded is not None:
    file_bytes = uploaded.getvalue()
    suffix = Path(uploaded.name).suffix.lower()

    # Initialize session state
    if "chat_history" not in st.session_state:
        st.session_state["chat_history"] = []
    if "chat_greeted" not in st.session_state:
        st.session_state["chat_greeted"] = False

    # Detect when a new file is uploaded (by content hash)
    file_key = hashlib.md5(file_bytes).hexdigest()
    if st.session_state.get("last_file_key") != file_key:
        st.session_state["last_file_key"] = file_key
        st.session_state.pop("annotated_text", None)
        st.session_state.pop("annotated_html", None)
        st.session_state["chat_history"] = []
        st.session_state["chat_greeted"] = False

    # If we haven't processed this document yet, extract text and annotate
    if st.session_state.get("annotated_html") is None:
        full_text = ""

        if suffix == ".pdf":
            if PdfReader is None:
                full_text = "(Install pypdf to enable PDF processing)"
            else:
                reader = PdfReader(io.BytesIO(file_bytes))
                parts = []
                for page in reader.pages:
                    parts.append(page.extract_text() or "")
                full_text = "\n".join(parts).strip()

        elif suffix == ".docx":
            if Document is None:
                full_text = "(Install python-docx to enable DOCX processing)"
            else:
                bio = io.BytesIO(file_bytes)
                doc = Document(bio)
                parts = [p.text for p in doc.paragraphs]
                full_text = "\n".join(parts).strip()

        else:
            full_text = "(Unsupported type)"

        # If we got real text, send it to the agent for placeholder cleaning & HTML
        if full_text and not full_text.startswith("("):
            status_ph = st.empty()
            progress = st.progress(0)
            messages = [
                "Analysing document…",
                "Detecting placeholders…",
                "Generating HTML preview…",
            ]

            result: dict = {}

            def _worker():
                ann = arun(clean_document_content(full_text))
                htmlp = arun(annotated_document_to_html(ann))
                result["annotated"] = ann
                result["html"] = htmlp

            t = threading.Thread(target=_worker, daemon=True)
            t.start()

            i = 0
            while t.is_alive():
                status_ph.info(messages[(i // 10) % len(messages)])
                progress.progress((i % 100) + 1)
                time.sleep(0.5)
                i = (i + 5) % 100

            t.join()
            progress.empty()
            status_ph.empty()

            annotated = result.get("annotated", full_text)
            html_page = result.get("html", "")

            st.session_state["annotated_text"] = annotated
            st.session_state["annotated_html"] = html_page
        else:
            # Just store the raw text / message
            st.session_state["annotated_text"] = full_text
            # For consistency, still go through the HTML renderer
            st.session_state["annotated_html"] = arun(
                annotated_document_to_html(full_text)
            )

    # Layout: left = chat, right = document preview & downloads
    left, right = st.columns([0.30, 0.70])

    # -------------------
    # Left: Chat column
    # -------------------
    with left:
        st.subheader("Chat")

        # Initial greeting once document is ready
        if (
            st.session_state["annotated_text"]
            and not st.session_state["chat_greeted"]
            and not st.session_state["annotated_text"].startswith("(")
        ):
            greeting = "Article processed, are you ready to fill it?"
            st.session_state["chat_history"].append(
                {"role": "assistant", "content": greeting}
            )
            st.session_state["chat_greeted"] = True

        user_msg = st.chat_input("Type to fill the form")
        if user_msg:
            st.session_state["chat_history"].append(
                {"role": "user", "content": user_msg}
            )

            with st.spinner("Updating form..."):
                res = arun(
                    run_form_chat_turn(
                        st.session_state["annotated_text"],
                        st.session_state["chat_history"],
                    )
                )

                for a in res.get("assistant_messages", []):
                    st.session_state["chat_history"].append(a)

                prev_annotated = st.session_state["annotated_text"]
                st.session_state["annotated_text"] = res.get(
                    "annotated", prev_annotated
                )

                if st.session_state["annotated_text"] != prev_annotated:
                    st.session_state["annotated_html"] = arun(
                        annotated_document_to_html(
                            st.session_state["annotated_text"]
                        )
                    )

        # Render chat history (newest at bottom)
        for m in reversed(st.session_state["chat_history"]):
            with st.chat_message(m["role"]):
                st.markdown(m["content"])

    # -------------------
    # Right: Preview column
    # -------------------
    with right:
        st.subheader("HTML Preview")
        st.components.v1.html(
            st.session_state.get("annotated_html", ""), height=900, scrolling=True
        )

        # Show download buttons when all placeholders are filled
        if st.session_state.get("annotated_text") and "<placeholder-start>" not in st.session_state[
            "annotated_text"
        ]:
            st.subheader("Download")
            final_text = annotated_to_plain(st.session_state["annotated_text"])

            # Build DOCX in-memory
            if Document is not None:
                bio = io.BytesIO()
                doc = Document()
                for line in final_text.splitlines():
                    doc.add_paragraph(_sanitize_for_docx(line))
                doc.save(bio)
                docx_bytes = bio.getvalue()

                st.download_button(
                    label="Download as DOCX",
                    data=docx_bytes,
                    file_name="completed_document.docx",
                    mime=(
                        "application/vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"
                    ),
                    use_container_width=True,
                )

                # Attempt PDF conversion via LibreOffice or docx2pdf if available
                try:
                    with tempfile.TemporaryDirectory() as td:
                        docx_path = Path(td) / "completed_document.docx"
                        with open(docx_path, "wb") as f:
                            f.write(docx_bytes)

                        pdf_path = convert_docx_to_pdf(str(docx_path))
                        with open(pdf_path, "rb") as f:
                            pdf_bytes = f.read()

                    st.download_button(
                        label="Download as PDF",
                        data=pdf_bytes,
                        file_name="completed_document.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                    )
                except Exception:
                    st.info(
                        "PDF conversion unavailable. Install 'docx2pdf' or LibreOffice "
                        "('soffice') to enable PDF export."
                    )
            else:
                st.info(
                    "DOCX export unavailable. Install python-docx to enable DOCX/PDF downloads."
                )

else:
    # No file uploaded yet – keep it quiet but avoid empty page
    st.info("Upload a PDF or DOCX document to get started.")
